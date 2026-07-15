"""
简历文件校验与文本提取服务。

服务先检查文件大小和扩展名，再把字节写入临时文件，复用 ``text_utils`` 的统一解析入口；
保留的格式专用方法提供 PDF、DOC、DOCX、TXT 的多级降级方案。无论解析成功与否，临时
文件都会在 ``finally`` 中清理。
"""
import logging
import os
import hashlib
import tempfile
from typing import BinaryIO, Tuple, Optional
from uuid import UUID
import mimetypes
from app.utils.text_utils import extract_text_content
from app.utils.file_utils import get_file_mime_type

logger = logging.getLogger(__name__)


class ResumeParserService:
    """校验简历文件元信息，并按 PDF、Word、文本格式提取统一正文。"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.doc', '.docx'}
    SUPPORTED_MIME_TYPES = {
        'application/pdf',
        'text/plain',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }

    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """只依据上传元数据执行轻量校验，返回“是否通过 + 可展示原因”。

        这里不读取文件内容，也不验证扩展名与真实 MIME 是否一致；调用方应在落盘和解析前
        显式调用本方法，真正无法解析的伪装文件仍会在后续提取阶段报错。
        """
        # 大小由调用方传入，因此此处只执行上限判断，不负责确认它与实际字节数一致。
        if file_size > self.max_file_size:
            return False, f"文件大小超过限制 ({self.max_file_size / 1024 / 1024}MB)"
        
        # 扩展名统一转为小写后判断，使 PDF、Pdf 等写法进入同一解析分支。
        _, ext = os.path.splitext(filename.lower())
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"不支持的文件格式。支持的格式: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        return True, "文件验证通过"

    async def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """把内存字节转为临时文件，并委托公共解析器提取纯文本。

        ``extract_text_content`` 接收文件路径而不是字节流，因此这里承担两种表示之间的转换。
        本方法不重复调用 ``validate_file``；文件类型和大小校验应由上游上传流程负责。
        """
        try:
            # MIME 由文件名推断，用于让公共解析器选择 PDF、Word 或纯文本读取策略。
            mime_type = get_file_mime_type(filename)

            # 临时文件保留原扩展名，因为部分第三方解析库会同时依赖路径后缀判断格式。
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name

            try:
                # 公共入口统一知识库上传与简历上传的格式处理和文本返回约定。
                extracted = await extract_text_content(temp_path, mime_type)
                return extracted
            finally:
                try:
                    # delete=False 允许 Windows 上的解析器重新打开文件，因此必须手动清理。
                    os.unlink(temp_path)
                except Exception:
                    # 清理失败不能覆盖已经得到的解析结果或原始解析异常。
                    pass
        except Exception as e:
            logger.error(f"提取文本内容失败: {e}")
            # 对外统一为“文件解析失败”，同时保留底层原因便于接口层展示或记录。
            raise Exception(f"文件解析失败: {str(e)}")

    async def _extract_from_txt(self, file_content: bytes) -> str:
        """按常见编码顺序把 TXT 字节解码为字符串。

        这是保留的格式专用兼容方法；当前公开入口优先委托 ``extract_text_content``。由于部分
        单字节编码几乎不会抛错，成功解码只表示字节合法，不保证文本语义一定正确。
        """
        try:
            # 从常见 Unicode/中文编码开始尝试，避免过早落入宽松的 latin-1。
            encodings = ['utf-8', 'utf-16', 'gbk', 'gb2312', 'big5', 'latin-1']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    # 当前编码不适用时继续尝试，不把一次解码失败视为整个文件失败。
                    continue
            
            # 最终兜底会丢弃非法字节，目标是尽可能返回可供后续 LLM 处理的文本。
            return file_content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            logger.error(f"TXT文件解析失败: {e}")
            raise

    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """从 PDF 各页提取文本并按换行拼接。

        优先使用 PyPDF2；只有未安装 PyPDF2 时才尝试 pdfplumber。若 PyPDF2 已安装但文件解析
        失败，会直接进入通用异常分支，并不会继续尝试 pdfplumber。
        """
        try:
            # PDF 阅读器直接消费内存字节流，不产生临时文件。
            import PyPDF2
            import io
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_content = []
            
            # 扫描件页面可能没有文本层，extract_text 返回空值时跳过该页。
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            
            return '\n'.join(text_content)
            
        except ImportError:
            # 该回退解决的是依赖缺失，不是对同一解析错误进行第二次尝试。
            try:
                import pdfplumber
                import io
                
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text_content = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                    
                    return '\n'.join(text_content)
                    
            except ImportError:
                logger.error("PDF解析库未安装，请安装 PyPDF2 或 pdfplumber")
                raise Exception("PDF解析功能不可用，请联系管理员")
                
        except Exception as e:
            logger.error(f"PDF文件解析失败: {e}")
            raise Exception(f"PDF文件解析失败: {str(e)}")

    async def _extract_from_doc(self, file_content: bytes) -> str:
        """按 textract、Word COM、antiword 的顺序解析旧版 DOC 文件。

        每种解析器都可能受操作系统或可选依赖限制，因此单个方案失败只记录警告并继续降级；
        全部失败后才向调用方抛错。原始 DOC 临时文件由内层 ``finally`` 统一删除。
        """
        try:
            # 多个解析器都只接受路径，先把上传字节落到唯一临时文件。
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # textract 跨平台且直接返回字节，优先尝试并按常见中文编码解码。
                try:
                    import textract
                    text_bytes = textract.process(temp_file_path)
                    # 第一个能完成解码且有内容的结果即作为标准字符串返回。
                    decoded = None
                    for enc in ['utf-8', 'utf-16', 'gbk', 'gb2312', 'big5', 'latin-1']:
                        try:
                            decoded = text_bytes.decode(enc)
                            break
                        except Exception:
                            continue
                    if decoded is None:
                        decoded = text_bytes.decode('utf-8', errors='ignore')
                    text = decoded.strip()
                    if text:
                        return text
                except ImportError:
                    logger.warning("未安装textract，尝试其他方式解析DOC")
                except Exception as e:
                    logger.warning(f"textract解析DOC失败: {e}")

                # COM 方案只适用于安装 Microsoft Word 和 pywin32 的 Windows 环境。
                try:
                    import win32com.client  # type: ignore
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(temp_file_path)
                    # Word 先另存为 TXT，再把文件内容转换回 Python 字符串。
                    temp_txt_path = temp_file_path + ".txt"
                    wdFormatText = 2
                    doc.SaveAs(temp_txt_path, FileFormat=wdFormatText)
                    doc.Close(False)
                    word.Quit()
                    # Word 输出编码受本机区域设置影响，因此仍需按候选编码读取。
                    text = None
                    for enc in ['utf-8', 'utf-16', 'gbk', 'gb2312', 'big5', 'latin-1']:
                        try:
                            with open(temp_txt_path, "r", encoding=enc) as f:
                                content = f.read().strip()
                                if content:
                                    text = content
                                    break
                        except UnicodeDecodeError:
                            continue
                        except Exception:
                            continue
                    if text is None:
                        with open(temp_txt_path, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read().strip()
                    os.unlink(temp_txt_path)
                    if text:
                        return text
                except ImportError:
                    logger.warning("未安装pywin32，跳过Word COM方式")
                except Exception as e:
                    logger.warning(f"Word COM方式解析DOC失败: {e}")

                # antiword 是最后的外部进程回退；退出码和非空标准输出必须同时满足才算成功。
                try:
                    import subprocess
                    result = subprocess.run(
                        ["antiword", "-m", "UTF-8", temp_file_path],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        check=False
                    )
                    # 即使已要求 antiword 输出 UTF-8，仍兼容不同系统包装层返回的编码。
                    text = None
                    for enc in ['utf-8', 'utf-16', 'gbk', 'gb2312', 'big5', 'latin-1']:
                        try:
                            text = result.stdout.decode(enc).strip()
                            if text:
                                break
                        except Exception:
                            continue
                    if text is None:
                        text = result.stdout.decode('utf-8', errors='ignore').strip()
                    if result.returncode == 0 and text:
                        return text
                    else:
                        logger.warning(f"antiword解析失败: {result.stderr.decode('utf-8', errors='ignore')}")
                except Exception as e:
                    logger.warning(f"调用antiword解析DOC失败: {e}")

                # 三种可选解析器都不可用或未返回有效文本时，才终止降级链。
                raise Exception("DOC文件解析失败：缺少可用的解析器，请安装textract或pywin32/antiword")

            finally:
                # 无论哪种解析器成功、失败或提前返回，都删除原始 DOC 临时文件。
                os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"DOC文件解析失败: {e}")
            raise Exception(f"DOC文件解析失败: {str(e)}")

    async def _extract_from_docx(self, file_content: bytes) -> str:
        """从 DOCX 内存流提取段落和表格单元格，并展平为换行文本。

        该方法保留文档中的可见文字，但不保留版式、图片及段落与表格的层级结构；当前公开
        解析入口通常通过公共 ``extract_text_content`` 完成同类工作。
        """
        try:
            from docx import Document
            import io
            
            # python-docx 可以直接读取内存流，不需要额外产生临时文件。
            doc = Document(io.BytesIO(file_content))
            
            # 空段落不会进入结果，降低后续 LLM 提示中的无效空白。
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 表格内容按“表 -> 行 -> 单元格”顺序追加到段落文本之后。
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content)
            
        except ImportError:
            logger.error("DOCX解析库未安装，请安装 python-docx")
            raise Exception("DOCX文件解析功能不可用，请联系管理员")
            
        except Exception as e:
            logger.error(f"DOCX文件解析失败: {e}")
            raise Exception(f"DOCX文件解析失败: {str(e)}")

    def get_file_info(self, filename: str, file_content: bytes) -> dict:
        """把原始上传字节整理为可存储的文件元数据字典。

        SHA-256 基于实际内容，可用于重复文件识别；``mime_type`` 仅由文件名猜测，可能为
        ``None``，不能作为文件真实性或安全性的证明。
        """
        file_size = len(file_content)
        file_hash = hashlib.sha256(file_content).hexdigest()
        _, ext = os.path.splitext(filename.lower())
        
        # 返回普通字符串和整数，便于直接写入 ORM 字段或序列化为 JSON。
        return {
            'filename': filename,
            'file_type': ext.lstrip('.'),
            'file_size': file_size,
            'file_hash': file_hash,
            'mime_type': mimetypes.guess_type(filename)[0]
        }